from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app_core import (
    APP_NAME,
    APP_VERSION,
    add_summary_to_packet,
    format_storage_gb,
    safe_filename,
    yes_no,
)
from settings_service import ensure_directories, get_output_paths


SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}


def add_heading(document, text, level=1):
    """
    Add a heading using Word's built-in heading styles.
    """
    document.add_heading(text, level=level)


def add_label_value_table(document, rows):
    """
    Add a simple two-column label/value table.
    """
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    for label, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(label)
        row_cells[1].text = "" if value is None else str(value)

    document.add_paragraph("")


def add_paragraph_block(document, text):
    """
    Add a paragraph block while preserving a blank value cleanly.
    """
    if text:
        for line in str(text).splitlines():
            document.add_paragraph(line)
    else:
        document.add_paragraph("")


def add_branding_image(document, settings):
    """
    Add optional patch/logo image at the top of the report.

    Supported now:
        PNG
        JPG/JPEG

    SVG/vector support is intentionally parked because DOCX image support is
    most reliable across Word versions with raster images.
    """
    branding = settings.get("report_branding", {})
    image_path_text = branding.get("patch_image_path", "").strip()

    if not image_path_text:
        return

    image_path = Path(image_path_text)

    if not image_path.exists():
        return

    if image_path.suffix.lower() not in SUPPORTED_IMAGE_EXTENSIONS:
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()

    try:
        run.add_picture(str(image_path), width=Inches(1.25))
    except Exception:
        return


def build_docx_report(packet, settings):
    """
    Build a DOCX acquisition packet report.
    """
    packet = add_summary_to_packet(packet)

    general = packet.get("general_info", {})
    intake = packet.get("intake_info", {})
    subject = packet.get("subject", {})
    processing = packet.get("processing", {})
    output = packet.get("output", {})
    report_info = packet.get("report_info", {})
    summary = packet.get("summary", {})
    media_summary = packet.get("media_examined_summary", {})
    department = packet.get("department", {})

    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    add_branding_image(document, settings)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Digital Evidence Acquisition Documentation")
    title_run.bold = True
    title_run.font.size = Pt(16)

    department_name = department.get("department_name", "")
    unit_name = department.get("unit_name", "")

    if department_name or unit_name:
        dept = document.add_paragraph()
        dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if department_name:
            dept.add_run(department_name)
        if department_name and unit_name:
            dept.add_run("\n")
        if unit_name:
            dept.add_run(unit_name)

    document.add_paragraph("")

    add_heading(document, "Administrative Information", level=1)
    add_label_value_table(document, [
        ("Case Number", general.get("case_number", "")),
        ("Agency Case Number", general.get("agency_case_number", "")),
        ("State / Local Case No.", general.get("state_local_case_number", "")),
        ("Case Type", general.get("case_type", "")),
        ("Offense / Incident", general.get("offense_or_incident", "")),
        ("City of Offense", general.get("city_of_offense", "")),
        ("State of Offense", general.get("state_of_offense", "")),
        ("Country of Offense", general.get("country_of_offense", "")),
        ("Subject", f"{subject.get('last_name', '')}, {subject.get('first_name', '')}"),
        ("Requesting Officer / Investigator", general.get("requesting_investigator", "")),
        ("Technician", general.get("technician", "")),
        ("Date Received", general.get("date_received", "")),
        ("Time Received", general.get("time_received", "")),
        ("Date Processed", general.get("date_processed", "")),
        ("Time Processed", general.get("time_processed", "")),
    ])

    add_heading(document, "Intake Information", level=1)
    add_label_value_table(document, [
        ("Drop-off Person", intake.get("dropoff_person", "")),
        ("Received From", intake.get("received_from", "")),
        ("Evidence Item Number", intake.get("evidence_item_number", "")),
        ("Checked Out From Evidence", yes_no(intake.get("checked_out_from_evidence", False))),
        ("Checked Out Date/Time", intake.get("checked_out_datetime", "")),
        ("Returned To Evidence", yes_no(intake.get("returned_to_evidence", False))),
        ("Returned Date/Time", intake.get("returned_datetime", "")),
        ("Evidence Location / Locker", intake.get("evidence_location", "")),
    ])

    add_heading(document, "Examination Information", level=1)
    add_label_value_table(document, [
        ("Exam Start Date", processing.get("exam_start_date", "")),
        ("Exam End Date", processing.get("exam_end_date", "")),
        ("Processing Type", processing.get("processing_type", "")),
        ("Processing Status", processing.get("processing_status", "")),
        ("Processing Notes", processing.get("processing_notes", "")),
    ])

    add_heading(document, "Device / Media Summary", level=1)
    add_label_value_table(document, [
        ("Total Devices / Media Count", summary.get("total_devices", 0)),
        ("Total Known Storage", format_storage_gb(summary.get("total_storage_gb"))),
        ("Total Media Examined", media_summary.get("total_media_examined", 0)),
        ("Hard Drive Credits", media_summary.get("hard_drive_credits", 0)),
        ("ETech Credits", media_summary.get("etech_credits", 0)),
        ("Media Credits", media_summary.get("media_credits", 0)),
    ])

    device_counts = summary.get("device_counts", {})
    if device_counts:
        add_heading(document, "Device Counts", level=2)
        add_label_value_table(
            document,
            [(device_type, count) for device_type, count in device_counts.items()]
        )

    add_heading(document, "Device / Media Detail", level=1)

    devices = packet.get("devices", [])

    if devices:
        for index, device in enumerate(devices, start=1):
            add_heading(document, f"Device Entry {index}", level=2)
            add_label_value_table(document, [
                ("Type", device.get("device_type", "")),
                ("Quantity", device.get("quantity", "")),
                ("Description", device.get("description", "")),
                ("Make", device.get("make", "")),
                ("Model", device.get("model", "")),
                ("Serial / Identifier", device.get("serial", "")),
                ("Storage Size Per Device", (
                    "Unknown"
                    if device.get("capacity_size") is None
                    else f"{device.get('capacity_size')} {device.get('capacity_unit', '')}"
                )),
                ("Total Known Storage", format_storage_gb(device.get("storage_total_gb"))),
                ("Volumes Examined", device.get("volumes_examined", "")),
                ("Volume Scale", device.get("volume_scale", "")),
                ("Encrypted", yes_no(device.get("encrypted", False))),
                ("Decrypted", yes_no(device.get("decrypted", False))),
                ("Tools Used to Decrypt", device.get("tools_used_to_decrypt", "")),
                ("Password Locked", yes_no(device.get("password_locked", False))),
                ("Password Unlocked", yes_no(device.get("password_unlocked", False))),
                ("Services Used to Unlock", device.get("services_used_to_unlock", "")),
            ])
    else:
        document.add_paragraph("No device/media entries recorded.")

    add_heading(document, "Tools Used", level=1)

    tools_used = packet.get("tools_used", [])

    if tools_used:
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "Tool"
        header_cells[1].text = "Version"

        for tool in tools_used:
            row_cells = table.add_row().cells
            row_cells[0].text = tool.get("name", "")
            row_cells[1].text = tool.get("version", "")
    else:
        document.add_paragraph("No tools entered.")

    document.add_paragraph("")

    add_heading(document, "Generated Output", level=1)
    add_label_value_table(document, [
        ("Output Type", output.get("output_type", "")),
        ("Output Filename / Identifier", output.get("output_filename", "")),
        ("Output Location", output.get("output_location", "")),
        ("Reader Report Generated", yes_no(output.get("reader_report_generated", False))),
        ("Case File Generated", yes_no(output.get("case_file_generated", False))),
    ])

    add_heading(document, "Other Data Analyzed", level=1)
    add_paragraph_block(document, report_info.get("other_data_analyzed", ""))

    add_heading(document, "Case Summary", level=1)
    add_paragraph_block(document, report_info.get("case_summary", ""))

    add_heading(document, "Technician Notes", level=1)
    add_paragraph_block(document, packet.get("technician_notes", ""))

    add_heading(document, "Limitations and Scope", level=1)
    add_paragraph_block(document, packet.get("scope_statement", ""))

    add_heading(document, "Technician Statement", level=1)
    document.add_paragraph(
        "I documented the above process based on the actions performed, tool output, "
        "and information available at the time of processing."
    )

    add_heading(document, "Report Generated", level=1)
    add_label_value_table(document, [
        ("Generated On", packet.get("created_at", "")),
        ("Generated By", f"{APP_NAME} v{APP_VERSION}"),
    ])

    return document


def save_docx_report(packet, settings=None):
    """
    Save the acquisition packet as a DOCX report.

    Returns:
        Path: path to DOCX report
    """
    if settings is None:
        settings = {}

    ensure_directories(settings)
    paths = get_output_paths(settings)

    packet = add_summary_to_packet(packet)

    general = packet.get("general_info", {})
    report_info = packet.get("report_info", {})

    case_number = general.get("case_number", "UNKNOWN_CASE")

    case_title = (
        general.get("offense_or_incident", "")
        or report_info.get("case_type", "")
        or general.get("case_type", "")
        or "Untitled_Case"
    )

    timestamp = packet.get("created_at", "") or "unknown_time"

    safe_case = safe_filename(case_number)
    safe_title = safe_filename(case_title)
    safe_timestamp = safe_filename(timestamp)

    base_filename = f"{safe_case}_{safe_title}_{safe_timestamp}_acquisition_packet"

    docx_path = paths["reports_dir"] / f"{base_filename}.docx"

    document = build_docx_report(packet, settings)
    document.save(docx_path)

    return docx_path